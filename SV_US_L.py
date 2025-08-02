# In[]
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import classification_report, confusion_matrix, accuracy_score, roc_auc_score, roc_curve
from docx import Document

# Daten laden
Soft_Version = "C:\\Users\\spast\\Desktop\\Learning Methods\\Data\\Fiktive_DataSet_SoftTrick.csv"
Image_SV = "C:\\Users\\spast\\Desktop\\Learning Methods\\Data\\Images Soft Trick\\"
Hard_Version = "C:\\Users\\spast\\Desktop\\Learning Methods\\Data\\Fiktive_DataSet_HardTrick.csv"
Image_HV = "C:\\Users\\spast\\Desktop\\Learning Methods\\Data\\Images Hard Trick\\"
mode = int(input("Soft Trick = 0 or Hard Trick = 1"))
if mode == 0:
    filename = Soft_Version
if mode == 1:
    filename = Hard_Version
data = pd.read_csv(filename)
#In[] CHECK DATA
counts = data["landed"].value_counts()
print(counts)
percentages = data["landed"].value_counts(normalize=True)*100
print(percentages)
import matplotlib.pyplot as plt

plt.rcParams["font.family"] = "Times New Roman"
data['landed'].value_counts().sort_index().plot(kind='bar', color=['red', 'green'])
plt.title('Frequency of landed', fontsize=18)
plt.xlabel('Value', fontsize=16)
plt.ylabel('Numbers', fontsize=16)
plt.xticks([0, 1], ['0 (not landed)', '1 (landed)'], rotation=0)
plt.grid(axis='y')
plt.tight_layout()
if "Soft" in filename:
    plt.savefig(Image_SV +"Distribution_Soft.png", dpi = 300)
else:
    plt.savefig(Image_HV + "Distribution_Hard.png", dpi = 300)
plt.show()

from scipy.stats import chisquare
# Beobachtete Werte
observed = data['landed'].value_counts().sort_index()
# Erwartete Werte (z. B. gleiche Wahrscheinlichkeit für 0 und 1)
expected = [len(data) / 2, len(data) / 2]
chi2, p = chisquare(f_obs=observed, f_exp=expected)
print("Chi2-Statistik:", chi2)
print("p-Wert:", p) # p > .05 unterscheidet sich nicht von einer 50/50 Verteilung
#In[]
# Schriftart & -größen setzen
plt.rcParams["font.family"] = "Times New Roman"
plt.rcParams["axes.titlesize"] = 18
plt.rcParams["axes.labelsize"] = 16
plt.rcParams["xtick.labelsize"] = 14
plt.rcParams["ytick.labelsize"] = 14
plt.rcParams["legend.fontsize"] = 14

# Spalten bereinigen
landed_col = data["landed"]
data = data.drop(columns=data.columns[:2])
data["landed"] = landed_col.astype("category")

# Numerische Spalten auswählen
num_cols = data.select_dtypes(include="number").columns

# Für jede Spalte einen Plot generieren und speichern
for col in num_cols:
    fig, ax = plt.subplots(figsize=(8, 5))

    sns.kdeplot(data=data, x=col, hue="landed", ax=ax, fill=True)
    ax.set_title(f"KDE from '{col}' to 'landed'")
    ax.set_xlabel(col)
    ax.set_ylabel("Density")

    plt.tight_layout()
    if "Soft" in filename:
        plt.savefig(Image_SV + f"kde_{col}.png", dpi=300)
    else:
        plt.savefig(Image_HV + f"kde_{col}.png", dpi=300)
    plt.close()
#In[] Features und Zielvariable definieren
X = data[["Age", "BodySize", "Experience", "Weight", "LegLength",
          "MentalRepresentation", "HipAbduction", "HipTorque_Nm"]]
y = data["landed"]  # 0 = Trick nicht gestanden, 1 = Trick gestanden

# Daten aufteilen
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=42, stratify=y) # 30 % des Testdatensatzes, stratify = Klassenverteilung in Trainings und Testdaten gleich bleibt

# Features skalieren (besonders wichtig für logistische Regression)
scaler = StandardScaler()
X_train_scaled = scaler.fit_transform(X_train)
X_test_scaled = scaler.transform(X_test)

# Modell trainieren
model = LogisticRegression()
model.fit(X_train_scaled, y_train)

# Vorhersagen
predictions = model.predict(X_test_scaled)
probs = model.predict_proba(X_test_scaled)[:, 1]  # Wahrscheinlichkeit für Klasse 1

# Evaluation
print("Confusion Matrix:\n", confusion_matrix(y_test, predictions))
print("\nClassification Report:\n", classification_report(y_test, predictions))
print("Accuracy:", accuracy_score(y_test, predictions))
print("ROC AUC:", roc_auc_score(y_test, probs))

plt.rcParams["font.family"] = "Times New Roman"
# ROC-Kurve plotten
fpr, tpr, _ = roc_curve(y_test, probs)
plt.figure(figsize=(6, 4))
plt.plot(fpr, tpr, label='ROC Curve (AUC = {:.2f})'.format(roc_auc_score(y_test, probs)))
plt.plot([0, 1], [0, 1], 'k--')  # Diagonale
plt.xlabel('False Positive Rate')
plt.ylabel('True Positive Rate')
plt.title('ROC Curve')
plt.legend()
plt.grid(True)
plt.tight_layout()

# Classification Report als Dictionary
report_dict = classification_report(y_test, predictions, output_dict=True)
# Neues Word-Dokument erstellen
doc = Document()
doc.add_heading('Classification Report', level=1)
# Tabelle erstellen (plus 1 für Kopfzeile)
table = doc.add_table(rows=1 + len(report_dict), cols=5)
table.style = 'Table Grid'
# Kopfzeile definieren
header_cells = table.rows[0].cells
header_cells[0].text = 'Class'
header_cells[1].text = 'Precision'
header_cells[2].text = 'Recall'
header_cells[3].text = 'F1-Score'
header_cells[4].text = 'Support'
# Datenzeilen einfügen
for i, (label, metrics) in enumerate(report_dict.items()):
    row_cells = table.rows[i + 1].cells
    row_cells[0].text = str(label)
    if isinstance(metrics, dict):  # Nur bei Klassen oder "macro avg", etc.
        row_cells[1].text = '{:.2f}'.format(metrics['precision'])
        row_cells[2].text = '{:.2f}'.format(metrics['recall'])
        row_cells[3].text = '{:.2f}'.format(metrics['f1-score'])
        row_cells[4].text = str(int(metrics['support']))
    else:
        row_cells[1].text = row_cells[2].text = row_cells[3].text = str(metrics)
        row_cells[4].text = '-'

if "Soft" in filename:
    plt.savefig("C:\\Users\\spast\\Desktop\\Learning Methods\\Results\\Soft Trick\\roc_curve_soft_trick.png", dpi=300, bbox_inches='tight')  # oder .pdf für Vektorformat
    # Word-Datei speichern
    doc.save("C:\\Users\\spast\\Desktop\\Learning Methods\\Results\\Soft Trick\\classification_report.docx")
else:
    plt.savefig("C:\\Users\\spast\\Desktop\\Learning Methods\\Results\\Hard Trick\\roc_curve_hard_trick.png", dpi=300, bbox_inches='tight')  # oder .pdf für Vektorformat
    doc.save("C:\\Users\\spast\\Desktop\\Learning Methods\\Results\\Hard Trick\\classification_report.docx")
#plt.show()

#In[]
def TestingNewAthlete(age = 22, bodysize = 178, experience = 3, weight = 70, leglength = 90,
                      mentalrepresentation = 4, hipabduction = 35, hiptorquenm = 130):
    new_athlete = {
    "Age": age,
    "BodySize": bodysize,
    "Experience": experience,
    "Weight": weight,
    "LegLength": leglength,
    "MentalRepresentation": mentalrepresentation,
    "HipAbduction": hipabduction,
    "HipTorque_Nm": hiptorquenm
    }
    new_data = pd.DataFrame([new_athlete])
    new_data_scaled = scaler.transform(new_data)
    predicted_class = model.predict(new_data_scaled)
    print("Vorhersage (0 = nicht gestanden, 1 = gestanden):", predicted_class[0])
    prob_success = model.predict_proba(new_data_scaled)[0][1]
    print("Wahrscheinlichkeit für Gelingen (Klasse 1): {:.2f}%".format(prob_success * 100))
    prozent = prob_success * 100
        # Word-Dokument erzeugen
    doc = Document()
    doc.add_heading("Ergebnis für neuen Athleten", level=1)
    
    # Eingabewerte auflisten
    doc.add_paragraph("Eingabewerte:")
    for key, value in new_athlete.items():
        doc.add_paragraph(f"{key}: {value}", style='List Bullet')
    
    # Ergebnistext
    text = f"\nEin Athlet mit den oben genannten Eigenschaften hat eine geschätzte Wahrscheinlichkeit von {prozent:.2f}% den Trick zu stehen."
    if predicted_class == 1:
        text += " Das Modell prognostiziert, dass der Athlet den Trick stehen wird."
    else:
        text += " Das Modell prognostiziert, dass der Athlet den Trick **nicht** stehen wird."
    doc.add_paragraph(text)
    
    # Dokument speichern
    if "Soft" in filename:
        doc.save("C:\\Users\\spast\\Desktop\\Learning Methods\\Results\\Soft Trick\\TestingNewAthlete_Report.docx")   
    if "Hard" in filename:
        doc.save("C:\\Users\\spast\\Desktop\\Learning Methods\\Results\\Hard Trick\\TestingNewAthlete_Report.docx")
    print("Word-Datei wurde gespeichert.")

TestingNewAthlete(age=30, bodysize=160, weight=79, experience=5, leglength=85,
                  mentalrepresentation=7, hipabduction=49, hiptorquenm=100)

#In[]
# Nach model.fit(X_train_scaled, y_train)

print("\nFeature Koeffizienten der logistischen Regression (standardisierte Features):")
for feature, coef in zip(X.columns, model.coef_[0]):
    print(f"{feature}: {coef:.4f}")

# Optional: Visualisierung der Koeffizienten
import matplotlib.pyplot as plt

coefs = model.coef_[0]
features = X.columns

plt.figure(figsize=(8,5))
plt.barh(features, coefs, color='teal')
plt.xlabel("Koeffizienten (Einfluss auf log-odds)")
plt.title("Feature Bedeutung aus der logistischen Regression")
plt.grid(axis='x')
plt.tight_layout()

# Grafik speichern je nach Modus

if "Soft" in filename:
    plt.savefig("C:\\Users\\spast\\Desktop\\Learning Methods\\Results\\Soft Trick\\" + "Feature_Importance_LogReg.png", dpi=300)
else:
    plt.savefig("C:\\Users\\spast\\Desktop\\Learning Methods\\Results\\Hard Trick\\" + "Feature_Importance_LogReg.png", dpi=300)
plt.show()
import numpy as np
import pandas as pd

# Odds Ratios berechnen
odds_ratios = np.exp(model.coef_[0])  # e^beta
odds_ratios_df = pd.DataFrame({
    'Feature': X.columns,
    'Odds Ratio': odds_ratios,
    'Koeffizient': model.coef_[0]
})

print("\nOdds Ratios (bezogen auf standardisierte Features):")
print(odds_ratios_df.sort_values("Odds Ratio", ascending=False))
plt.figure(figsize=(8, 5))
sorted_df = odds_ratios_df.sort_values("Odds Ratio", ascending=True)

plt.barh(sorted_df["Feature"], sorted_df["Odds Ratio"], color='teal')
plt.xscale("log")  # log-Skala für bessere Lesbarkeit
plt.xlabel("Odds Ratio (log-skaliert)")
plt.title("Einfluss der Features auf die Wahrscheinlichkeit (Odds Ratio)")
plt.grid(True, which="both", axis="x", linestyle="--")
plt.tight_layout()

# Grafik speichern
if "Soft" in filename:
    plt.savefig("C:\\Users\\spast\\Desktop\\Learning Methods\\Results\\Soft Trick\\" + "Odds_Ratios_LogReg.png", dpi=300)
else:
    plt.savefig("C:\\Users\\spast\\Desktop\\Learning Methods\\Results\\Hard Trick\\" + "Odds_Ratios_LogReg.png", dpi=300)

plt.show()
#In[] P-Werte
import statsmodels.api as sm
# X und y definieren
X = data[["Age", "Experience", "BodySize", "HipAbduction", "Weight", 
          "MentalRepresentation", "LegLength", "HipTorque_Nm"]]
X = sm.add_constant(X)  # Intercept hinzufügen
y = data["landed"].astype(int)  # y sollte integer sein

# Modell fitten
model = sm.Logit(y, X)
result = model.fit()

# Zusammenfassung anzeigen
print(result.summary())
# Zusammenfassung in DataFrame
summary_table = pd.DataFrame({
    "Feature": result.params.index,
    "Koeffizient (β)": result.params.values,
    "Odds Ratio (e^β)": np.exp(result.params.values),
    "p-Wert": result.pvalues,
    "95% CI Untergrenze": np.exp(result.conf_int()[0]),
    "95% CI Obergrenze": np.exp(result.conf_int()[1])
})

# Sortiert nach kleinsten p-Werten (bedeutungsvollste Features oben)
summary_table = summary_table.sort_values("p-Wert")

# Ausgabe
pd.set_option('display.float_format', '{:.4f}'.format)
print(summary_table)
# Neues Word-Dokument erstellen
doc = Document()
doc.add_heading('Logistische Regression: Zusammenfassung der Features', level=1)

# Tabelle erstellen
table = doc.add_table(rows=1, cols=len(summary_table.columns))
table.style = 'Table Grid'

# Kopfzeile einfügen
hdr_cells = table.rows[0].cells
for i, col_name in enumerate(summary_table.columns):
    hdr_cells[i].text = str(col_name)

# Datenzeilen einfügen
for _, row in summary_table.iterrows():
    row_cells = table.add_row().cells
    for i, val in enumerate(row):
        if isinstance(val, float):
            row_cells[i].text = f"{val:.4f}"
        else:
            row_cells[i].text = str(val)

# Speichern – je nach Trick-Modus in richtigen Ordner
if "Soft" in filename:
    output_path = "C:\\Users\\spast\\Desktop\\Learning Methods\\Results\\Soft Trick\\feature_significance_logit.docx"
else:
    output_path = "C:\\Users\\spast\\Desktop\\Learning Methods\\Results\\Hard Trick\\feature_significance_logit.docx"

doc.save(output_path)
print("Word-Tabelle wurde gespeichert unter:", output_path)

'''
Die ROC-Kurve (Receiver Operating Characteristic) ist ein Evaluierungswerkzeug 
für binäre Klassifikationsmodelle, das zeigt, wie gut dein Modell zwischen den zwei Klassen 
(z.B. 0 = nicht gelandet, 1 = gelandet) unterscheidet, bei allen möglichen Entscheidungsschwellen.

TP (True Positive):
Das Modell sagt „gelandet“ (1) - und das stimmt.
→ Richtig positiv erkannt.

FP (False Positive):
Das Modell sagt „gelandet“ (1) - aber in Wahrheit ist es nicht gelandet (0).
→ Falsch-Alarm.

FN (False Negative):
Das Modell sagt „nicht gelandet“ (0) - aber in Wirklichkeit wurde doch gelandet (1).
→ Verpasster Treffer.

TN (True Negative):
Das Modell sagt „nicht gelandet“ (0) - und das ist korrekt.
→ Richtig negativ erkannt.

Was du in der Kurve siehst:
Jeder Punkt auf der Kurve entspricht einer anderen Threshold-Einstellung (also z.B. 
„Ab wann ist ein Wert wahrscheinlich genug für Klasse 1?“).

Die diagonale Linie (von unten links nach oben rechts) zeigt reines Raten (Chance = 50%).

Je weiter die Kurve nach oben links gebogen ist, desto besser ist dein Modell.

Der Flächeninhalt unter der ROC-Kurve (AUC) fasst die Performance in einer Zahl zwischen 0 und 1 zusammen.

AUC = 1.0 → perfekter Klassifikator

AUC = 0.5 → wie Zufall

AUC < 0.5 → schlimmer als Raten (wahrscheinlich invertierte Vorhersage)'''
#In[] UNSUPERVISED LEARNING
# Imports und Daten laden
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns

from sklearn.preprocessing import StandardScaler
from sklearn.decomposition import PCA
from sklearn.cluster import KMeans

# Daten laden
Soft_Version = "C:\\Users\\spast\\Desktop\\Learning Methods\\Data\\Fiktive_DataSet_SoftTrick.csv"
Image_SV = "C:\\Users\\spast\\Desktop\\Learning Methods\\Data\\Images Soft Trick\\"
Hard_Version = "C:\\Users\\spast\\Desktop\\Learning Methods\\Data\\Fiktive_DataSet_HardTrick.csv"
Image_HV = "C:\\Users\\spast\\Desktop\\Learning Methods\\Data\\Images Hard Trick\\"

filename = Soft_Version
data = pd.read_csv(filename)

# Nur die Features (ohne Zielvariable)
X = data[["Age", "BodySize", "Experience", "Weight", "LegLength",
          "MentalRepresentation", "HipAbduction", "HipTorque_Nm"]]

# Standardisierung alle Features haben den gleichen Mittelwert und Standardabweichung
scaler = StandardScaler() # Z transformation über den Mittelwert und Standardabweichung berechnen (X - Mittelwert)/Standardabweichung
X_scaled = scaler.fit_transform(X) # sensitiv gegenüber unterschiedlicher Skalierung

# Hauptkomponentenanalyse PCA durchführen, PCA: welche Features zur Varianz beitragen
pca = PCA(n_components=None)
pca.fit(X_scaled)
plt.rcParams["font.family"] = "Times New Roman"
# Kumulative erklärte Varianz plotten
plt.figure(figsize=(6, 4))
plt.plot(np.cumsum(pca.explained_variance_ratio_), marker='o')
plt.xlabel("Anzahl der Komponenten")
plt.ylabel("Kumulative erklärte Varianz")
plt.title("PCA - erklärte Varianz")
plt.grid(True)
plt.tight_layout()
if "Hard" in filename:
    plt.savefig("C:\\Users\\spast\\Desktop\\Learning Methods\\Results\\Hard Trick\\Kumulative Variance_hard_trick.png", dpi=300, bbox_inches='tight')
if "Soft" in filename:
    plt.savefig("C:\\Users\\spast\\Desktop\\Learning Methods\\Results\\Soft Trick\\Kumulative Variance_soft_trick.png", dpi=300, bbox_inches='tight')  
plt.show()

# Komponenten-Loadings anzeigen (wichtigste Features pro Komponente)
loadings = pd.DataFrame(pca.components_.T, index=X.columns, # enthält die Koeefizienten der Originalfeatures
                        columns=[f"PC{i+1}" for i in range(len(X.columns))])
print("\nWichtigste Features für PC1 (nach Einfluss sortiert):")
print(loadings["PC1"].sort_values(ascending=False)) #Tabelle: welche Werte am meisten zur Hauptkomponente zu beitragen

# KMeans-Clustering Gruppenaufteilung
# K random points, euclidean distance between the chosen cluster points and the other points
kmeans = KMeans(n_clusters=2, random_state=42) # Kann KMEANS zwei echte Gruppen landed(not landed) herausfiltern
clusters = kmeans.fit_predict(X_scaled) # after calculating the distance from the chosen to each point, recalculating the center of each cluster and redefining each points' distance
data["Cluster"] = clusters

# PCA-Reduktion auf 2D für Plot
pca_2d = PCA(n_components=2)
X_pca = pca_2d.fit_transform(X_scaled)

# Visualisierung der Cluster, PC1 Richtung mit der größten Streeung, PC2 zweitgrößte Streuung 
plt.figure(figsize=(6, 4))
plt.scatter(X_pca[:, 0], X_pca[:, 1], c=clusters, cmap="viridis", s=50)
plt.xlabel("PC1")
plt.ylabel("PC2")
plt.title("KMeans-Cluster (visualisiert über PCA)")
plt.grid(True)
plt.tight_layout()
if "Hard" in filename:
    plt.savefig("C:\\Users\\spast\\Desktop\\Learning Methods\\Results\\Hard Trick\\Cluster_hard_trick.png", dpi=300, bbox_inches='tight')
if "Soft" in filename:
    plt.savefig("C:\\Users\\spast\\Desktop\\Learning Methods\\Results\\Soft Trick\\Cluster_soft_trick.png", dpi=300, bbox_inches='tight')  
plt.show()

# Cluster-Mittelwerte der Features
print("\nDurchschnittswerte der Features pro Cluster:")
cluster_means = data.groupby("Cluster")[X.columns].mean().T
print(cluster_means)

# Optional: Heatmap für bessere Übersicht
plt.figure(figsize=(8, 4))
sns.heatmap(cluster_means, annot=True, cmap="coolwarm")
plt.title("Cluster-Mittelwerte der Features")
plt.tight_layout()
if "Hard" in filename:
    plt.savefig("C:\\Users\\spast\\Desktop\\Learning Methods\\Results\\Hard Trick\\Heatmap_hard_trick.png", dpi=300, bbox_inches='tight')
if "Soft" in filename:
    plt.savefig("C:\\Users\\spast\\Desktop\\Learning Methods\\Results\\Soft Trick\\Heatmap_soft_trick.png", dpi=300, bbox_inches='tight') 
plt.show()

pd.crosstab(data["landed"], data["Cluster"])
# Falls nötig: Cluster-Labels umkehren
data["Cluster_corrected"] = np.where(data["Cluster"] == 1, 1, 0)
from sklearn.metrics import accuracy_score
print("Accuracy zwischen Clustering und gelabelten Daten:")
print(accuracy_score(data["landed"], data["Cluster_corrected"]))
# %%